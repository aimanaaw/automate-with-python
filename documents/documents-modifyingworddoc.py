import docx

documentObject = docx.Document('demo.docx')
print(documentObject)
print(documentObject.paragraphs[0].text)

p = documentObject.paragraphs[1]
print(p.runs)
# Each change in text is a run
# A plain paragraph having some bold and some italic. => 4 runs in this object
print(p.runs[0].text)
# p.runs[1].bold = True

p.runs[3].underline = True
p.runs[3].text = 'italic and underlined'
documentObject.save('demo2.docx')

p.style = 'Title'
documentObject.save('demo3.docx')

d = docx.Document()
d.add_paragraph('Hello this is a paragraph')
d.add_paragraph('Hello this is another paragraph')
d.save('demo4.docx')
p1 = d.paragraphs[0]
p1.add_run('\nThis is a new run')
p1.runs[1].bold = True
print(p1.runs[1])
d.save('demo6.docx')
